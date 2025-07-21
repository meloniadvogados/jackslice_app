import streamlit as st
import os
import tempfile
import zipfile
import subprocess
from io import BytesIO
import base64

def clear_converter_data():
    """
    Limpa todos os dados do converter e força reset completo
    """
    # Limpar dados específicos do converter
    keys_to_clear = [
        'converter_pdf_data', 'converter_filename', 'converter_total_pages',
        'converter_file_size_mb', 'converter_results', 'converter_pdf_with_ocr'
    ]
    
    for key in keys_to_clear:
        if key in st.session_state:
            del st.session_state[key]
    
    # Incrementar key do uploader para forçar reset
    st.session_state.converter_uploader_key += 1

def pdf_converter_page():
    """
    Página para converter PDFs para vários formatos
    """
    st.title("🪄 Converter PDF")

    # Gerar uma key única para o uploader baseada em um contador
    if 'converter_uploader_key' not in st.session_state:
        st.session_state.converter_uploader_key = 0
    
    # Upload do PDF
    uploaded_file = st.file_uploader(
        "📄 Selecione um arquivo PDF",
        type=['pdf'],
        help="Formatos aceitos: PDF",
        key=f"converter_uploader_{st.session_state.converter_uploader_key}"
    )
    
    if uploaded_file is not None:
        # Armazenar dados do PDF
        pdf_data = uploaded_file.getvalue()
        
        # Obter informações do PDF e verificar se tem texto
        try:
            import PyPDF2
            from io import BytesIO
            
            pdf_reader = PyPDF2.PdfReader(BytesIO(pdf_data))
            total_pages = len(pdf_reader.pages)
            file_size_mb = len(pdf_data) / (1024 * 1024)
            
            # Verificar se PDF tem texto pesquisável
            has_text = check_pdf_has_text(pdf_data)
            needs_ocr = not has_text
            
            # Exibir informações do arquivo
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("📄 Total de Páginas", total_pages)
            with col2:
                if file_size_mb < 0.1:
                    size_display = f"{file_size_mb:.3f} MB"
                else:
                    size_display = f"{file_size_mb:.1f} MB"
                st.metric("📦 Tamanho", size_display)
            with col3:
                avg_size = file_size_mb / total_pages if total_pages > 0 else 0
                st.metric("📊 Média/Página", f"{avg_size:.2f} MB")
            
            # Status de OCR
            col1, col2 = st.columns(2)
            with col1:
                if has_text:
                    st.success("✅ **PDF com texto pesquisável** - Pronto para conversão!")
                else:
                    st.warning("⚠️ **PDF digitalizado** - Requer OCR para melhor conversão")
            
            with col2:
                if needs_ocr:
                    st.info("🧠 **OCR será aplicado automaticamente** durante a conversão")
                else:
                    st.info("✅ OCR não necessário")
            
            st.markdown("---")
            
            # Usar PDF com OCR se disponível
            final_pdf_data = st.session_state.get('converter_pdf_with_ocr', pdf_data)
            
            # Tabs para diferentes categorias de conversão
            tab1, tab2, tab3 = st.tabs(["📄 Documentos", "🖼️ Imagens", "🌐 Web & Outros"])
            
            with tab1:
                convert_to_documents(final_pdf_data, total_pages, uploaded_file.name)
            
            with tab2:
                convert_to_images(final_pdf_data, total_pages, uploaded_file.name)
            
            with tab3:
                convert_to_web_others(final_pdf_data, total_pages, uploaded_file.name)
                
        except Exception as e:
            st.error(f"❌ Erro ao processar PDF: {str(e)}")
    
    else:
        # Informações sobre conversão quando não há arquivo carregado
        with st.expander("ℹ️ Informações sobre Conversão de PDF"):
            st.markdown("""
            **Como funciona a conversão:**
            
            Esta ferramenta converte PDFs para diferentes formatos utilizando:
            - **pdf2docx**: Conversão para Word mantendo layout original
            - **pdfplumber**: Extração de texto e tabelas
            - **PyMuPDF**: Extração de imagens em alta qualidade
            - **openpyxl**: Criação de planilhas Excel
            
            **Formatos suportados:**
            
            📄 **Documentos:**
            - **Word (.docx)**: Layout idêntico ao PDF + texto editável
            - **Excel (.xlsx)**: Tabelas e dados estruturados
            - **PowerPoint (.pptx)**: Cada página como slide (em desenvolvimento)
            - **Texto (.txt)**: Texto puro extraído
            
            🖼️ **Imagens:**
            - **PNG**: Melhor qualidade, suporte a transparência
            - **JPEG**: Menor tamanho, ideal para fotos
            - **TIFF**: Qualidade profissional
            
            🌐 **Web & Outros:**
            - **HTML**: Páginas web navegáveis (em desenvolvimento)
            - **ePub**: E-books compatíveis com leitores digitais (em desenvolvimento)
            - **XML**: Dados estruturados (em desenvolvimento)
            
            **Recursos especiais:**
            - Detecção automática de texto pesquisável
            - **OCR automático** para PDFs digitalizados
            - Múltiplas opções de qualidade e configuração
            - Seleção de páginas específicas
            """)
            
            st.markdown("**💡 Dica:** O OCR é aplicado automaticamente quando necessário - sem etapas manuais!")


def convert_to_documents(pdf_data, total_pages, filename):
    """
    Conversão para formatos de documentos (Word, Excel, PowerPoint, TXT)
    """
    st.markdown("### 📄 Converter para Documentos")
    st.markdown("💡 **Ideal para**: Editar texto, trabalhar com planilhas, criar apresentações")
    
    # Seleção do formato
    format_choice = st.radio(
        "**Escolha o formato de documento:**",
        ["📝 Word (.docx)", "📊 Excel (.xlsx)", "🎯 PowerPoint (.pptx)", "📋 Texto (.txt)"],
        help="Selecione o formato que melhor se adequa ao seu objetivo"
    )
    
    # Opções específicas para cada formato
    if format_choice == "📝 Word (.docx)":
        st.markdown("**📝 Exemplo**: Perfeito para editar contratos, relatórios e documentos de texto")
        st.markdown("**🎯 Use quando**: Precisar editar o conteúdo do PDF")
        
        # Opções avançadas
        with st.expander("⚙️ Opções Avançadas"):
            preserve_layout = st.checkbox("Preservar layout original", value=True, 
                                        help="Mantém formatação similar ao PDF original")
            include_images = st.checkbox("Incluir imagens", value=True,
                                       help="Converte imagens do PDF para o Word")
            
            # Mostrar opção de método apenas se houver texto pesquisável
            final_pdf_data = st.session_state.get('converter_pdf_with_ocr', pdf_data)
            if check_pdf_has_text(final_pdf_data):
                conversion_method = st.radio(
                    "Método de conversão:",
                    ["🚀 HÍBRIDO: Imagens + Texto editável (RECOMENDADO)", 
                     "🎨 Layout visual: Preserva aparência original"],
                    help="Híbrido = cabeçalhos/logos preservados + texto editável | Layout = apenas aparência visual"
                )
                use_text_method = "HÍBRIDO" in conversion_method
            else:
                use_text_method = False
        
        st.markdown("---")
        col1, col2 = st.columns(2)
        with col1:
            if st.button("🪄 Converter para Word", type="primary", use_container_width=True, key="convert_word"):
                convert_to_word(final_pdf_data, filename, preserve_layout, include_images, use_text_method)
        with col2:
            if st.button("🔄 Começar Novamente", type="secondary", use_container_width=True, key="restart_word"):
                clear_converter_data()
                st.rerun()
    
    elif format_choice == "📊 Excel (.xlsx)":
        st.markdown("**📊 Exemplo**: Ideal para PDFs com tabelas, dados financeiros e planilhas")
        st.markdown("**🎯 Use quando**: O PDF contém tabelas que você quer editar")
        
        # Opções específicas
        with st.expander("⚙️ Opções de Tabela"):
            detect_tables = st.checkbox("Detectar tabelas automaticamente", value=True,
                                      help="Procura e converte tabelas automaticamente")
            merge_cells = st.checkbox("Manter células mescladas", value=False,
                                    help="Preserva formatação de células mescladas")
        
        st.markdown("---")
        col1, col2 = st.columns(2)
        with col1:
            if st.button("🪄 Converter para Excel", type="primary", use_container_width=True, key="convert_excel"):
                convert_to_excel(pdf_data, filename, detect_tables, merge_cells)
        with col2:
            if st.button("🔄 Começar Novamente", type="secondary", use_container_width=True, key="restart_excel"):
                clear_converter_data()
                st.rerun()
    
    elif format_choice == "🎯 PowerPoint (.pptx)":
        st.markdown("**🎯 Exemplo**: Transforma cada página do PDF em um slide")
        st.markdown("**🎯 Use quando**: Quiser usar o PDF como base para apresentação")
        
        # Opções específicas
        with st.expander("⚙️ Opções de Apresentação"):
            slides_per_page = st.number_input("Páginas por slide", min_value=1, max_value=4, value=1,
                                            help="Quantas páginas do PDF por slide")
            add_notes = st.checkbox("Adicionar notas aos slides", value=False,
                                  help="Extrai texto para as notas do slide")
        
        st.markdown("---")
        col1, col2 = st.columns(2)
        with col1:
            if st.button("🪄 Converter para PowerPoint", type="primary", use_container_width=True, key="convert_ppt"):
                convert_to_powerpoint(pdf_data, filename, slides_per_page, add_notes)
        with col2:
            if st.button("🔄 Começar Novamente", type="secondary", use_container_width=True, key="restart_ppt"):
                clear_converter_data()
                st.rerun()
    
    else:  # Texto
        st.markdown("**📋 Exemplo**: Extrai todo o texto do PDF em formato simples")
        st.markdown("**🎯 Use quando**: Quiser apenas o texto, sem formatação")
        
        # Opções específicas
        with st.expander("⚙️ Opções de Texto"):
            preserve_layout = st.checkbox("Preservar quebras de linha", value=True,
                                        help="Mantém estrutura de parágrafos")
            include_page_numbers = st.checkbox("Incluir números de página", value=False,
                                             help="Adiciona indicador de página no texto")
        
        st.markdown("---")
        col1, col2 = st.columns(2)
        with col1:
            if st.button("🪄 Converter para Texto", type="primary", use_container_width=True, key="convert_txt"):
                convert_to_text(pdf_data, filename, preserve_layout, include_page_numbers)
        with col2:
            if st.button("🔄 Começar Novamente", type="secondary", use_container_width=True, key="restart_txt"):
                clear_converter_data()
                st.rerun()


def convert_to_images(pdf_data, total_pages, filename):
    """
    Conversão para formatos de imagem (PNG, JPEG, TIFF)
    """
    st.markdown("### 🖼️ Converter para Imagens")
    st.markdown("💡 **Ideal para**: Criar imagens de alta qualidade, usar em apresentações, sites")
    
    # Seleção do formato
    format_choice = st.radio(
        "**Escolha o formato de imagem:**",
        ["🖼️ PNG (melhor qualidade)", "📸 JPEG (menor tamanho)", "🎨 TIFF (profissional)"],
        help="PNG = transparência e qualidade | JPEG = tamanho menor | TIFF = uso profissional"
    )
    
    # Configurações de qualidade
    with st.expander("⚙️ Configurações de Qualidade"):
        col1, col2 = st.columns(2)
        
        with col1:
            dpi = st.selectbox(
                "Resolução (DPI):",
                [150, 300, 600, 1200],
                index=1,
                help="150=Web | 300=Impressão | 600=Alta qualidade | 1200=Profissional"
            )
        
        with col2:
            if "JPEG" in format_choice:
                quality = st.slider("Qualidade JPEG:", 50, 100, 90, 
                                  help="Maior valor = melhor qualidade e maior tamanho")
            else:
                quality = 90  # Default para outros formatos
    
    # Opções de páginas
    with st.expander("📄 Seleção de Páginas"):
        page_option = st.radio(
            "Quais páginas converter?",
            ["Todas as páginas", "Páginas específicas", "Intervalo de páginas"]
        )
        
        if page_option == "Páginas específicas":
            pages_input = st.text_input(
                "Digite as páginas (ex: 1,3,5):",
                placeholder="1,3,5,10",
                help="Use vírgulas para separar as páginas"
            )
        elif page_option == "Intervalo de páginas":
            col1, col2 = st.columns(2)
            with col1:
                start_page = st.number_input("Página inicial:", min_value=1, max_value=total_pages, value=1)
            with col2:
                end_page = st.number_input("Página final:", min_value=start_page, max_value=total_pages, value=total_pages)
        else:
            pages_input = None
    
    # Mostrar estimativa
    if page_option == "Todas as páginas":
        estimated_files = total_pages
    elif page_option == "Páginas específicas" and pages_input:
        try:
            pages = [int(p.strip()) for p in pages_input.split(',') if p.strip()]
            estimated_files = len(pages)
        except:
            estimated_files = 0
    elif page_option == "Intervalo de páginas":
        estimated_files = end_page - start_page + 1
    else:
        estimated_files = total_pages
    
    if estimated_files > 0:
        st.success(f"✅ **Resultado**: {estimated_files} imagem(ns) será(ão) gerada(s)")
    
    st.markdown("---")
    col1, col2 = st.columns(2)
    with col1:
        if st.button("🪄 Converter para Imagens", type="primary", use_container_width=True, key="convert_images"):
            # Preparar parâmetros
            if page_option == "Páginas específicas" and pages_input:
                try:
                    page_list = [int(p.strip()) for p in pages_input.split(',') if p.strip()]
                except:
                    st.error("❌ Formato de páginas inválido")
                    return
            elif page_option == "Intervalo de páginas":
                page_list = list(range(start_page, end_page + 1))
            else:
                page_list = None
            
            convert_to_image_format(pdf_data, filename, format_choice, dpi, quality, page_list)
    with col2:
        if st.button("🔄 Começar Novamente", type="secondary", use_container_width=True, key="restart_images"):
            clear_converter_data()
            st.rerun()


def convert_to_web_others(pdf_data, total_pages, filename):
    """
    Conversão para web e outros formatos (HTML, ePub, XML)
    """
    st.markdown("### 🌐 Converter para Web & Outros")
    st.markdown("💡 **Ideal para**: Publicação web, e-books, dados estruturados")
    
    # Seleção do formato
    format_choice = st.radio(
        "**Escolha o formato:**",
        ["🌐 HTML (páginas web)", "📚 ePub (e-book)", "📋 XML (dados estruturados)"],
        help="HTML = para sites | ePub = para leitores digitais | XML = dados estruturados"
    )
    
    if format_choice == "🌐 HTML (páginas web)":
        st.markdown("**🌐 Exemplo**: Cria páginas web que podem ser publicadas online")
        st.markdown("**🎯 Use quando**: Quiser publicar o conteúdo em um site")
        
        with st.expander("⚙️ Opções de HTML"):
            include_css = st.checkbox("Incluir estilos CSS", value=True,
                                    help="Adiciona formatação visual básica")
            embed_images = st.checkbox("Incorporar imagens", value=True,
                                     help="Inclui imagens diretamente no HTML")
            single_page = st.checkbox("Página única", value=False,
                                    help="Cria um único arquivo HTML ao invés de múltiplos")
        
        st.markdown("---")
        col1, col2 = st.columns(2)
        with col1:
            if st.button("🪄 Converter para HTML", type="primary", use_container_width=True, key="convert_html"):
                convert_to_html(pdf_data, filename, include_css, embed_images, single_page)
        with col2:
            if st.button("🔄 Começar Novamente", type="secondary", use_container_width=True, key="restart_html"):
                clear_converter_data()
                st.rerun()
    
    elif format_choice == "📚 ePub (e-book)":
        st.markdown("**📚 Exemplo**: Cria um e-book compatível com Kindle, Kobo, etc.")
        st.markdown("**🎯 Use quando**: Quiser ler o PDF em dispositivos móveis")
        
        with st.expander("⚙️ Opções de ePub"):
            book_title = st.text_input("Título do livro:", value=filename.replace('.pdf', ''))
            book_author = st.text_input("Autor:", value="Convertido por Jack PDF Slicer")
            add_toc = st.checkbox("Adicionar índice", value=True,
                                help="Cria um índice navegável no e-book")
        
        st.markdown("---")
        col1, col2 = st.columns(2)
        with col1:
            if st.button("🪄 Converter para ePub", type="primary", use_container_width=True, key="convert_epub"):
                convert_to_epub(pdf_data, filename, book_title, book_author, add_toc)
        with col2:
            if st.button("🔄 Começar Novamente", type="secondary", use_container_width=True, key="restart_epub"):
                clear_converter_data()
                st.rerun()
    
    else:  # XML
        st.markdown("**📋 Exemplo**: Extrai dados estruturados para processamento automático")
        st.markdown("**🎯 Use quando**: Precisar dos dados em formato estruturado")
        
        with st.expander("⚙️ Opções de XML"):
            include_metadata = st.checkbox("Incluir metadados", value=True,
                                         help="Adiciona informações sobre páginas e formatação")
            extract_tables = st.checkbox("Extrair tabelas", value=True,
                                       help="Identifica e estrutura tabelas separadamente")
            preserve_structure = st.checkbox("Preservar estrutura", value=True,
                                           help="Mantém hierarquia de parágrafos e seções")
        
        st.markdown("---")
        col1, col2 = st.columns(2)
        with col1:
            if st.button("🪄 Converter para XML", type="primary", use_container_width=True, key="convert_xml"):
                convert_to_xml(pdf_data, filename, include_metadata, extract_tables, preserve_structure)
        with col2:
            if st.button("🔄 Começar Novamente", type="secondary", use_container_width=True, key="restart_xml"):
                clear_converter_data()
                st.rerun()


# Funções de conversão específicas

def convert_to_word(pdf_data, filename, preserve_layout, include_images, use_text_method=None):
    """Converte PDF para Word - método híbrido para melhor qualidade"""
    try:
        with st.spinner("🪄 Convertendo para Word..."):
            # Verificar se precisa de OCR
            if not check_pdf_has_text(pdf_data):
                st.info("🧠 PDF digitalizado detectado - Aplicando OCR automaticamente...")
                
                # Aplicar OCR automaticamente
                ocr_pdf_data = apply_ocr_to_pdf(pdf_data, filename)
                if ocr_pdf_data:
                    st.success("✅ OCR aplicado com sucesso!")
                    pdf_data = ocr_pdf_data  # Usar PDF com OCR
                else:
                    st.warning("⚠️ Falha no OCR - Convertendo PDF original")
            
            # Sempre usar pdf2docx para preservar layout EXATAMENTE IGUAL
            convert_to_word_pdf2docx(pdf_data, filename, preserve_layout, include_images)
            
    except Exception as e:
        st.error(f"❌ Erro na conversão para Word: {str(e)}")


def convert_to_word_text_based(pdf_data, filename, preserve_layout):
    """Converte PDF para Word HÍBRIDO: imagens + texto editável (melhor para PDFs com OCR)"""
    try:
        import pdfplumber
        import fitz  # PyMuPDF para extrair imagens
        from docx import Document
        from docx.shared import Inches, Cm
        from io import BytesIO
        from PIL import Image
        
        # Criar documento Word
        doc = Document()
        
        # Adicionar título
        title = doc.add_heading(filename.replace('.pdf', ''), 0)
        
        # Abrir PDF com ambas as bibliotecas
        pdf_fitz = fitz.open(stream=pdf_data, filetype="pdf")
        total_pages = len(pdf_fitz)
        
        # Criar barra de progresso
        progress_bar = st.progress(0, text="🔄 Iniciando conversão híbrida...")
        
        with pdfplumber.open(BytesIO(pdf_data)) as pdf_plumber:
            for page_num, (page_plumber, page_fitz) in enumerate(zip(pdf_plumber.pages, pdf_fitz), 1):
                # Adicionar separador de página (exceto primeira)
                if page_num > 1:
                    doc.add_page_break()
                
                # === 1. EXTRAIR E ADICIONAR IMAGENS ESPECÍFICAS DA PÁGINA ===
                try:
                    # Extrair imagens específicas (gráficos, fotos) do PDF
                    image_list = page_fitz.get_images()
                    
                    for img_index, img in enumerate(image_list):
                        try:
                            # Extrair dados da imagem
                            xref = img[0]
                            base_image = page_fitz.extract_image(xref)
                            image_bytes = base_image["image"]
                            
                            # Adicionar imagem ao Word
                            img_stream = BytesIO(image_bytes)
                            paragraph = doc.add_paragraph()
                            run = paragraph.add_run()
                            
                            # Redimensionar imagem (largura máxima 12cm)
                            run.add_picture(img_stream, width=Cm(12))
                            doc.add_paragraph()  # Espaço após imagem
                            
                        except Exception as img_err:
                            # Continuar se houver erro com uma imagem específica
                            continue
                    
                except Exception as e:
                    # Continuar se houver erro geral com imagens
                    pass
                
                # === 2. EXTRAIR E ADICIONAR TEXTO EDITÁVEL ===
                progress_bar.progress(page_num / total_pages, text=f"🔄 Processando página {page_num} de {total_pages}")
                
                # Adicionar separador visual
                doc.add_paragraph("─" * 50)
                doc.add_paragraph(f"📝 TEXTO EDITÁVEL - Página {page_num}").bold = True
                doc.add_paragraph("─" * 50)
                
                # Extrair texto
                text = page_plumber.extract_text()
                if text and text.strip():
                    if preserve_layout:
                        # Preservar quebras de linha
                        paragraphs = text.split('\n\n')
                        for para_text in paragraphs:
                            if para_text.strip():
                                p = doc.add_paragraph(para_text.strip())
                    else:
                        # Texto corrido
                        clean_text = ' '.join(text.split())
                        if clean_text.strip():
                            doc.add_paragraph(clean_text)
                else:
                    doc.add_paragraph("(Nenhum texto editável encontrado nesta página)")
                
                # === 3. EXTRAIR E ADICIONAR TABELAS ===
                tables = page_plumber.extract_tables()
                if tables:
                    doc.add_paragraph()
                    doc.add_paragraph("📊 TABELAS DETECTADAS").bold = True
                    
                    for table_idx, table_data in enumerate(tables):
                        if table_data and len(table_data) > 0:
                            try:
                                # Verificar se a tabela tem dados válidos
                                max_cols = max(len(row) for row in table_data if row)
                                if max_cols > 0:
                                    # Adicionar tabela ao Word
                                    word_table = doc.add_table(rows=len(table_data), cols=max_cols)
                                    word_table.style = 'Table Grid'
                                    
                                    for row_idx, row in enumerate(table_data):
                                        for col_idx in range(max_cols):
                                            cell_value = row[col_idx] if col_idx < len(row) else ""
                                            if cell_value:
                                                word_table.cell(row_idx, col_idx).text = str(cell_value)
                                    
                                    doc.add_paragraph()  # Espaço após tabela
                            except Exception as table_err:
                                st.error(f"Erro ao processar tabela {table_idx} da página {page_num}: {table_err}")
                
                doc.add_paragraph()  # Espaço entre páginas
        
        # Fechar documento PyMuPDF
        pdf_fitz.close()
        
        # Finalizar barra de progresso
        progress_bar.progress(1.0, text="✅ Conversão híbrida concluída!")
        
        # Salvar documento Word em BytesIO
        docx_buffer = BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        
        # Download
        base_name = filename.replace('.pdf', '') if filename.endswith('.pdf') else filename
        output_filename = f"{base_name}_HIBRIDO.docx"
        st.success("✅ Conversão HÍBRIDA para Word concluída!")
        st.info("📋 **Resultado**: Imagens preservadas + Texto editável + Tabelas estruturadas")
        
        # Verificar se o buffer tem conteúdo
        docx_data = docx_buffer.getvalue()
        if len(docx_data) == 0:
            st.error("❌ Erro: Documento Word vazio gerado!")
            return
        
        
        st.download_button(
            label="📥 Baixar Word Híbrido",
            data=docx_data,
            file_name=output_filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary"
        )
        
    except Exception as e:
        st.error(f"❌ Erro na conversão híbrida: {str(e)}")
        st.error("💡 Tentando método alternativo...")
        
        # Fallback para método simples de texto
        try:
            convert_to_word_simple_text(pdf_data, filename, preserve_layout)
        except Exception as e2:
            st.error(f"❌ Erro no método alternativo: {str(e2)}")


def convert_to_word_simple_text(pdf_data, filename, preserve_layout):
    """Método simplificado apenas com texto (fallback)"""
    try:
        import pdfplumber
        from docx import Document
        from io import BytesIO
        
        doc = Document()
        title = doc.add_heading(filename.replace('.pdf', ''), 0)
        
        with pdfplumber.open(BytesIO(pdf_data)) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                if page_num > 1:
                    doc.add_page_break()
                
                doc.add_heading(f'Página {page_num}', level=2)
                
                text = page.extract_text()
                if text:
                    if preserve_layout:
                        paragraphs = text.split('\n\n')
                        for para_text in paragraphs:
                            if para_text.strip():
                                doc.add_paragraph(para_text.strip())
                    else:
                        clean_text = ' '.join(text.split())
                        doc.add_paragraph(clean_text)
        
        docx_buffer = BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        
        output_filename = filename.replace('.pdf', '_TEXTO.docx')
        st.warning("⚠️ Conversão simplificada (apenas texto)")
        st.download_button(
            label="📥 Baixar Word Simplificado",
            data=docx_buffer.getvalue(),
            file_name=output_filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="secondary"
        )
        
    except Exception as e:
        st.error(f"❌ Erro no método simplificado: {str(e)}")


def convert_to_word_pdf2docx(pdf_data, filename, preserve_layout, include_images):
    """Converte PDF para Word usando pdf2docx (melhor para PDFs nativos)"""
    try:
        from pdf2docx import Converter
        
        # Criar arquivos temporários
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_pdf:
            tmp_pdf.write(pdf_data)
            tmp_pdf_path = tmp_pdf.name
        
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_docx:
            tmp_docx_path = tmp_docx.name
        
        # Converter
        cv = Converter(tmp_pdf_path)
        cv.convert(tmp_docx_path)
        cv.close()
        
        # Verificar se o arquivo foi criado
        if not os.path.exists(tmp_docx_path):
            st.error("❌ Arquivo Word não foi criado!")
            return
        
        # Ler arquivo convertido
        with open(tmp_docx_path, 'rb') as f:
            docx_data = f.read()
        
        # Verificar se o arquivo tem conteúdo
        if len(docx_data) == 0:
            st.error("❌ Arquivo Word está vazio!")
            return
        
        # Limpar arquivos temporários
        os.unlink(tmp_pdf_path)
        os.unlink(tmp_docx_path)
        
        # Download
        base_name = filename.replace('.pdf', '') if filename.endswith('.pdf') else filename
        output_filename = f"{base_name}.docx"
        
        st.success("✅ Conversão para Word concluída!")
        st.info(f"📄 **Arquivo gerado**: {output_filename} ({len(docx_data)/1024/1024:.1f} MB)")
        st.warning("⚠️ **Importante**: Se o arquivo baixar como PDF, renomeie a extensão para .docx")
        
        st.download_button(
            label="📥 Baixar arquivo Word",
            data=docx_data,
            file_name=output_filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary"
        )
        
    except Exception as e:
        st.error(f"❌ Erro na conversão pdf2docx: {str(e)}")


def convert_to_text(pdf_data, filename, preserve_layout, include_page_numbers):
    """Converte PDF para texto usando pdfplumber"""
    try:
        with st.spinner("🪄 Extraindo texto..."):
            # Verificar se precisa de OCR
            if not check_pdf_has_text(pdf_data):
                st.info("🧠 PDF digitalizado detectado - Aplicando OCR automaticamente...")
                
                # Aplicar OCR automaticamente
                ocr_pdf_data = apply_ocr_to_pdf(pdf_data, filename)
                if ocr_pdf_data:
                    st.success("✅ OCR aplicado com sucesso!")
                    pdf_data = ocr_pdf_data  # Usar PDF com OCR
                else:
                    st.warning("⚠️ Falha no OCR - Convertendo PDF original")
            import pdfplumber
            from io import BytesIO
            
            text_content = []
            
            with pdfplumber.open(BytesIO(pdf_data)) as pdf:
                for i, page in enumerate(pdf.pages, 1):
                    if include_page_numbers:
                        text_content.append(f"\n--- Página {i} ---\n")
                    
                    page_text = page.extract_text()
                    if page_text:
                        if preserve_layout:
                            text_content.append(page_text)
                        else:
                            # Remover quebras de linha excessivas
                            clean_text = ' '.join(page_text.split())
                            text_content.append(clean_text)
                    
                    text_content.append("\n")
            
            final_text = '\n'.join(text_content)
            
            # Verificar se há conteúdo
            if not final_text.strip():
                st.error("❌ Nenhum texto foi extraído do PDF!")
                return
            
            # Download
            base_name = filename.replace('.pdf', '') if filename.endswith('.pdf') else filename
            output_filename = f"{base_name}.txt"
            text_data = final_text.encode('utf-8')
            
            st.success("✅ Conversão para texto concluída!")
            st.info(f"📝 **Arquivo gerado**: {output_filename} ({len(text_data)/1024:.1f} KB)")
            st.info(f"📊 **Caracteres extraídos**: {len(final_text):,}")
            
            st.download_button(
                label="📥 Baixar arquivo de texto",
                data=text_data,
                file_name=output_filename,
                mime="text/plain",
                type="primary"
            )
            
    except Exception as e:
        st.error(f"❌ Erro na conversão para texto: {str(e)}")


def convert_to_image_format(pdf_data, filename, format_choice, dpi, quality, page_list):
    """Converte PDF para imagens usando PyMuPDF"""
    try:
        with st.spinner("🪄 Convertendo para imagens..."):
            import fitz
            
            # Determinar formato
            if "PNG" in format_choice:
                img_format = "png"
                file_ext = "png"
            elif "JPEG" in format_choice:
                img_format = "jpeg"
                file_ext = "jpg"
            else:  # TIFF
                img_format = "tiff"
                file_ext = "tiff"
            
            doc = fitz.open(stream=pdf_data, filetype="pdf")
            
            # Determinar páginas a converter
            if page_list:
                pages_to_convert = [p-1 for p in page_list if 0 <= p-1 < doc.page_count]
            else:
                pages_to_convert = range(doc.page_count)
            
            if len(pages_to_convert) == 1:
                # Uma única imagem
                page = doc[pages_to_convert[0]]
                mat = fitz.Matrix(dpi/72, dpi/72)
                pix = page.get_pixmap(matrix=mat)
                
                if img_format == "jpeg":
                    img_data = pix.jpeg_image(quality=quality)
                elif img_format == "png":
                    img_data = pix.png_image()
                else:  # TIFF
                    img_data = pix.pil_tobytes(format="TIFF")
                
                output_filename = f"{filename.replace('.pdf', '')}_page_{pages_to_convert[0]+1}.{file_ext}"
                
                st.success("✅ Conversão para imagem concluída!")
                st.download_button(
                    label="📥 Baixar imagem",
                    data=img_data,
                    file_name=output_filename,
                    mime=f"image/{img_format}",
                    type="primary"
                )
            else:
                # Múltiplas imagens em ZIP
                zip_buffer = BytesIO()
                with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                    
                    for page_num in pages_to_convert:
                        page = doc[page_num]
                        mat = fitz.Matrix(dpi/72, dpi/72)
                        pix = page.get_pixmap(matrix=mat)
                        
                        if img_format == "jpeg":
                            img_data = pix.jpeg_image(quality=quality)
                        elif img_format == "png":
                            img_data = pix.png_image()
                        else:  # TIFF
                            img_data = pix.pil_tobytes(format="TIFF")
                        
                        img_filename = f"page_{page_num+1}.{file_ext}"
                        zip_file.writestr(img_filename, img_data)
                
                zip_buffer.seek(0)
                output_filename = f"{filename.replace('.pdf', '')}_images.zip"
                
                st.success(f"✅ {len(pages_to_convert)} imagens convertidas!")
                st.download_button(
                    label="📥 Baixar imagens (ZIP)",
                    data=zip_buffer.getvalue(),
                    file_name=output_filename,
                    mime="application/zip",
                    type="primary"
                )
            
            doc.close()
            
    except Exception as e:
        st.error(f"❌ Erro na conversão para imagens: {str(e)}")


def convert_to_excel(pdf_data, filename, detect_tables, merge_cells):
    """Converte PDF para Excel (implementação básica)"""
    try:
        with st.spinner("🪄 Convertendo para Excel..."):
            # Verificar se precisa de OCR
            if not check_pdf_has_text(pdf_data):
                st.info("🧠 PDF digitalizado detectado - Aplicando OCR automaticamente...")
                
                # Aplicar OCR automaticamente
                ocr_pdf_data = apply_ocr_to_pdf(pdf_data, filename)
                if ocr_pdf_data:
                    st.success("✅ OCR aplicado com sucesso!")
                    pdf_data = ocr_pdf_data  # Usar PDF com OCR
                else:
                    st.warning("⚠️ Falha no OCR - Convertendo PDF original")
            import pdfplumber
            import openpyxl
            from io import BytesIO
            
            # Criar workbook
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "PDF Content"
            
            current_row = 1
            
            with pdfplumber.open(BytesIO(pdf_data)) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    # Adicionar separador de página
                    ws.cell(row=current_row, column=1, value=f"=== Página {page_num} ===")
                    current_row += 2
                    
                    if detect_tables:
                        # Tentar extrair tabelas
                        tables = page.extract_tables()
                        if tables:
                            for table in tables:
                                for row_data in table:
                                    for col_num, cell_value in enumerate(row_data, 1):
                                        if cell_value:
                                            ws.cell(row=current_row, column=col_num, value=str(cell_value))
                                    current_row += 1
                                current_row += 1  # Espaço entre tabelas
                        else:
                            # Se não há tabelas, extrair texto
                            text = page.extract_text()
                            if text:
                                for line in text.split('\n'):
                                    if line.strip():
                                        ws.cell(row=current_row, column=1, value=line.strip())
                                        current_row += 1
                    else:
                        # Extrair apenas texto
                        text = page.extract_text()
                        if text:
                            for line in text.split('\n'):
                                if line.strip():
                                    ws.cell(row=current_row, column=1, value=line.strip())
                                    current_row += 1
                    
                    current_row += 2  # Espaço entre páginas
            
            # Salvar em BytesIO
            excel_buffer = BytesIO()
            wb.save(excel_buffer)
            excel_buffer.seek(0)
            
            # Verificar se o arquivo tem conteúdo
            excel_data = excel_buffer.getvalue()
            if len(excel_data) == 0:
                st.error("❌ Arquivo Excel está vazio!")
                return
            
            # Download
            base_name = filename.replace('.pdf', '') if filename.endswith('.pdf') else filename
            output_filename = f"{base_name}.xlsx"
            
            st.success("✅ Conversão para Excel concluída!")
            st.info(f"📊 **Arquivo gerado**: {output_filename} ({len(excel_data)/1024/1024:.1f} MB)")
            st.warning("⚠️ **Importante**: Se o arquivo baixar como PDF, renomeie a extensão para .xlsx")
            
            st.download_button(
                label="📥 Baixar arquivo Excel",
                data=excel_data,
                file_name=output_filename,
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                type="primary"
            )
            
    except Exception as e:
        st.error(f"❌ Erro na conversão para Excel: {str(e)}")


# Implementações básicas para outros formatos
def convert_to_powerpoint(pdf_data, filename, slides_per_page, add_notes):
    """Conversão para PowerPoint (implementação futura)"""
    st.warning("🚧 Conversão para PowerPoint em desenvolvimento!")
    st.info("💡 Em breve: Cada página do PDF se tornará um slide")

def convert_to_html(pdf_data, filename, include_css, embed_images, single_page):
    """Conversão para HTML (implementação futura)"""
    st.warning("🚧 Conversão para HTML em desenvolvimento!")
    st.info("💡 Em breve: PDF como páginas web navegáveis")

def convert_to_epub(pdf_data, filename, book_title, book_author, add_toc):
    """Conversão para ePub (implementação futura)"""
    st.warning("🚧 Conversão para ePub em desenvolvimento!")
    st.info("💡 Em breve: E-books compatíveis com todos os leitores")

def convert_to_xml(pdf_data, filename, include_metadata, extract_tables, preserve_structure):
    """Conversão para XML (implementação futura)"""
    st.warning("🚧 Conversão para XML em desenvolvimento!")
    st.info("💡 Em breve: Dados estruturados para processamento automático")


# Funções auxiliares para OCR

def check_pdf_has_text(pdf_data):
    """Verifica se o PDF já tem texto pesquisável"""
    try:
        import PyPDF2
        from io import BytesIO
        
        pdf_reader = PyPDF2.PdfReader(BytesIO(pdf_data))
        
        # Verificar algumas páginas em busca de texto
        pages_to_check = min(3, len(pdf_reader.pages))  # Verificar até 3 páginas
        total_text_length = 0
        
        for i in range(pages_to_check):
            try:
                page_text = pdf_reader.pages[i].extract_text()
                if page_text:
                    total_text_length += len(page_text.strip())
            except:
                continue
        
        # Se encontrou texto significativo (mais de 50 caracteres), considera que tem texto
        return total_text_length > 50
        
    except Exception as e:
        print(f"Erro ao verificar texto do PDF: {e}")
        return False  # Em caso de erro, assume que precisa de OCR


def apply_ocr_to_pdf(pdf_data, filename):
    """Aplica OCR no PDF usando a lógica do módulo OCR existente"""
    try:
        # Salvar arquivo temporário de entrada
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_input:
            tmp_input.write(pdf_data)
            input_path = tmp_input.name
        
        # Arquivo de saída temporário
        with tempfile.NamedTemporaryFile(delete=False, suffix='_ocr.pdf') as tmp_output:
            output_path = tmp_output.name
        
        # Mostrar progresso
        with st.spinner("🧠 Aplicando OCR... Isso pode levar alguns minutos..."):
            # Comando OCRmyPDF com configurações básicas
            command = [
                "ocrmypdf",
                "--language", "por",  # Português por padrão
                "--jobs", "2",
                "--jpeg-quality", "85",
                "--output-type", "pdf",
                "--force-ocr",  # Força OCR mesmo se já tiver texto
                input_path,
                output_path
            ]
            
            try:
                # Executar comando
                result = subprocess.run(
                    command,
                    capture_output=True,
                    text=True,
                    timeout=300  # 5 minutos timeout
                )
                
                if result.returncode == 0:
                    # Ler arquivo processado
                    with open(output_path, 'rb') as f:
                        ocr_pdf_data = f.read()
                    
                    # Limpar arquivos temporários
                    os.unlink(input_path)
                    os.unlink(output_path)
                    
                    return ocr_pdf_data
                else:
                    st.error(f"❌ Erro no OCR: {result.stderr}")
                    return None
                    
            except subprocess.TimeoutExpired:
                st.error("❌ OCR demorou muito para processar. Tente com um PDF menor.")
                return None
            except FileNotFoundError:
                st.error("❌ OCRmyPDF não encontrado. Verifique se está instalado.")
                return None
            except Exception as e:
                st.error(f"❌ Erro inesperado no OCR: {str(e)}")
                return None
        
    except Exception as e:
        st.error(f"❌ Erro ao preparar OCR: {str(e)}")
        return None
    finally:
        # Garantir limpeza dos arquivos temporários
        try:
            if 'input_path' in locals() and os.path.exists(input_path):
                os.unlink(input_path)
            if 'output_path' in locals() and os.path.exists(output_path):
                os.unlink(output_path)
        except:
            pass